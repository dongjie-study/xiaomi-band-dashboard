"""生成618复盘分析总结Word文档（以手表销量为核心）"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

# ========== 读取数据 ==========
with open(r'C:\Users\Administrator\Desktop\小米手环直播间销量分析\618_analysis_data.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

# ========== 工具函数 ==========
def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_border(table):
    """给表格设置边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        borders.append(border)
    tblPr.append(borders)

def add_styled_table(doc, headers, rows, col_widths=None, highlight_header=True):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if highlight_header:
            set_cell_shading(cell, '1E90FF')

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F0F4F8')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_title(doc, text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    return h

def add_para(doc, text, bold=False, color=None, size=10.5, space_after=6):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, bold_prefix="", size=10.5):
    """添加带符号的列表项"""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(size)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

# ========== 读取数据 ==========
teams = data['teams']
our = teams['我方']
comp = teams['良米']
jixie = teams['机械空间']

# 手表数据
watch6_total = data['prod_REDMI Watch 6']
watch_s_total = data['prod_小米手表 S系列']

our_products = data['our_products']
comp_products = data['competitor_products']

our_watch6 = our_products.get('REDMI Watch 6', {})
our_watch_s = our_products.get('小米手表 S系列', {})
comp_watch6 = comp_products.get('REDMI Watch 6', {})
comp_watch_s = comp_products.get('小米手表 S系列', {})

# ========== 生成文档 ==========
doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ===== 封面标题 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
run = p.add_run('618大促复盘分析总结')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('——以手表品类销量为核心')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('分析周期：2026年5月15日 — 6月18日（35天）')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数据来源：抖音直播间+商品卡订单数据、千川视频投放数据、违规监控数据')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()

# ========================================
# 第一部分：618整体概况
# ========================================
add_title(doc, '一、618大促整体概况', level=1)

add_para(doc, '1.1 核心数据总览', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_styled_table(doc,
    ['指标', '数值', '说明'],
    [
        ['全周期总订单', '165,405单', '5/15-6/18，共35天'],
        ['全周期总GSV', '6,073.5万元', ''],
        ['直播间订单', '133,188单（80.5%）', '主力销售渠道'],
        ['商品卡订单', '32,217单（19.5%）', '货架电商增量'],
        ['618当天总订单', '8,845单', '全品类峰值'],
        ['我方618当天订单', '3,441单（38.9%）', '6/1的1.7倍'],
    ],
    col_widths=[4.5, 5, 6]
)

doc.add_paragraph()
add_para(doc, '1.2 各团队表现对比', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_styled_table(doc,
    ['团队', '订单数', 'GSV（万元）', '份额', '直播间数', '单号均产'],
    [
        ['我方', f'{our["orders"]:,}', f'{our["gsv_wan"]}', f'{our["pct"]}%', str(our["rooms"]), f'{our["avg_per_room"]:,}'],
        ['良米', f'{comp["orders"]:,}', f'{comp["gsv_wan"]}', f'{comp["pct"]}%', str(comp["rooms"]), f'{comp["avg_per_room"]:,}'],
        ['机械空间', f'{jixie["orders"]:,}', f'{jixie["gsv_wan"]}', f'{jixie["pct"]}%', str(jixie["rooms"]), f'{jixie["avg_per_room"]:,}'],
        ['综训', f'{teams["综训"]["orders"]:,}', f'{teams["综训"]["gsv_wan"]}', f'{teams["综训"]["pct"]}%', str(teams["综训"]["rooms"]), f'{teams["综训"]["avg_per_room"]:,}'],
    ],
    col_widths=[2.5, 2.5, 3, 2, 2.5, 3]
)

doc.add_paragraph()

# ========================================
# 第二部分：手表品类深度分析（核心）
# ========================================
add_title(doc, '二、手表品类销量深度分析（核心）', level=1)

add_para(doc, '2.1 手表品类整体表现', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_para(doc, '手表品类是本次618中GSV贡献最高的品类，REDMI Watch 6以2,020万元GSV超越小米手环10（1,648万元）成为全品类GSV冠军。但我方在手表品类的份额严重落后于竞对良米。')

watch_total_orders = watch6_total['total'] + watch_s_total['total']
watch_total_gsv = watch6_total['total_gsv'] + watch_s_total['total_gsv']
our_watch_orders = our_watch6.get('orders', 0) + our_watch_s.get('orders', 0)
comp_watch_orders = comp_watch6.get('orders', 0) + comp_watch_s.get('orders', 0)

add_styled_table(doc,
    ['产品', '全品类订单', '全品类GSV（元）', '我方订单', '竞对订单', '我方份额'],
    [
        ['REDMI Watch 6', f'{watch6_total["total"]:,}', f'¥{watch6_total["total_gsv"]:,.0f}',
         f'{our_watch6.get("orders",0):,}', f'{comp_watch6.get("orders",0):,}', f'{watch6_total["share"]}%'],
        ['小米手表 S系列', f'{watch_s_total["total"]:,}', f'¥{watch_s_total["total_gsv"]:,.0f}',
         f'{our_watch_s.get("orders",0):,}', f'{comp_watch_s.get("orders",0):,}', f'{watch_s_total["share"]}%'],
        ['手表合计', f'{watch_total_orders:,}', f'¥{watch_total_gsv:,.0f}',
         f'{our_watch_orders:,}', f'{comp_watch_orders:,}', f'{our_watch_orders/watch_total_orders*100:.1f}%'],
    ],
    col_widths=[3.5, 2.5, 3.5, 2.5, 2.5, 2]
)

doc.add_paragraph()
add_para(doc, '2.2 REDMI Watch 6——我方与竞对全面对比', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_para(doc, 'REDMI Watch 6是本次618中竞争最激烈的单品，全品类44,766单、2,020万元GSV。我方仅占29.8%份额，差距达18,114单。')

add_styled_table(doc,
    ['对比维度', '我方', '竞对（良米）', '差距', '分析'],
    [
        ['总订单', f'{our_watch6.get("orders",0):,}', f'{comp_watch6.get("orders",0):,}',
         f'-{watch6_total["diff"]:,}', '竞对是我方的2.36倍'],
        ['总GSV', f'¥{our_watch6.get("our_gsv",0):,.0f}', f'¥{comp_watch6.get("comp_gsv",0):,.0f}',
         '', '竞对GSV是我方的2.36倍'],
        ['直播间订单', f'{our_watch6.get("our_live",0):,}', f'{comp_watch6.get("comp_live",0):,}',
         '', '直播间是主战场'],
        ['商品卡订单', f'{our_watch6.get("our_card",0):,}', f'{comp_watch6.get("comp_card",0):,}',
         '', '商品卡差距相对小'],
        ['直播间GSV', f'¥{our_watch6.get("our_live_gsv",0):,.0f}', f'¥{comp_watch6.get("comp_live_gsv",0):,.0f}',
         '', '直播间贡献主要GSV'],
        ['商品卡GSV', f'¥{our_watch6.get("our_card_gsv",0):,.0f}', f'¥{comp_watch6.get("comp_card_gsv",0):,.0f}',
         '', '商品卡有提升空间'],
        ['市场份额', '29.8%', '70.2%', '-40.4pp', '差距巨大'],
    ],
    col_widths=[3, 3, 3, 2.5, 4.5]
)

doc.add_paragraph()
add_para(doc, '2.3 小米手表 S系列——高端手表严重缺失', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_para(doc, '小米手表 S系列客单价高（均价约1,000元），是利润最高的手表产品，但我方份额仅11.3%，严重落后于竞对。')

add_styled_table(doc,
    ['对比维度', '我方', '竞对（良米）', '差距'],
    [
        ['总订单', f'{our_watch_s.get("orders",0):,}', f'{comp_watch_s.get("orders",0):,}', f'-{watch_s_total["diff"]:,}'],
        ['总GSV', f'¥{our_watch_s.get("our_gsv",0):,.0f}', f'¥{comp_watch_s.get("comp_gsv",0):,.0f}', ''],
        ['直播间订单', f'{our_watch_s.get("our_live",0):,}', f'{comp_watch_s.get("comp_live",0):,}', ''],
        ['商品卡订单', f'{our_watch_s.get("our_card",0):,}', f'{comp_watch_s.get("comp_card",0):,}', ''],
        ['市场份额', '11.3%', '88.7%', '-77.4pp'],
    ],
    col_widths=[3, 3.5, 3.5, 3]
)

doc.add_paragraph()
add_para(doc, '2.4 手表品类直播间矩阵对比', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_para(doc, '竞对在手表品类拥有5个专属直播间，合计贡献24,021单；我方仅2个手表号，合计10,731单。直播间数量差距是手表品类落后的重要原因。')

# 我方手表直播间
add_para(doc, '我方手表直播间：', bold=True, size=10.5)
our_stores = data['我方_store_list']
watch_rooms_our = [
    ('小米官方手表', our_stores.get('小米官方手表', {}).get('orders', 0), our_stores.get('小米官方手表', {}).get('gsv', 0)),
    ('小米官旗手表直播间', our_stores.get('小米官旗手表直播间', {}).get('orders', 0), our_stores.get('小米官旗手表直播间', {}).get('gsv', 0)),
]
add_styled_table(doc,
    ['直播间', '订单数', 'GSV（元）'],
    [[r[0], f'{r[1]:,}', f'¥{r[2]:,.0f}'] for r in watch_rooms_our] + [['合计', f'{sum(r[1] for r in watch_rooms_our):,}', f'¥{sum(r[2] for r in watch_rooms_our):,.0f}']],
    col_widths=[5, 3.5, 4]
)

doc.add_paragraph()
# 竞对手表直播间
add_para(doc, '竞对（良米）手表直播间：', bold=True, size=10.5)
comp_stores = data['良米_store_list']
watch_rooms_comp = [
    ('小米手表官方直播间', comp_stores.get('小米手表官方直播间', {}).get('orders', 0), comp_stores.get('小米手表官方直播间', {}).get('gsv', 0)),
    ('小米手表', comp_stores.get('小米手表', {}).get('orders', 0), comp_stores.get('小米手表', {}).get('gsv', 0)),
    ('小米智能手表旗舰店', comp_stores.get('小米智能手表旗舰店', {}).get('orders', 0), comp_stores.get('小米智能手表旗舰店', {}).get('gsv', 0)),
    ('小米手表官旗直播间', comp_stores.get('小米手表官旗直播间', {}).get('orders', 0), comp_stores.get('小米手表官旗直播间', {}).get('gsv', 0)),
    ('小米手表直播间', comp_stores.get('小米手表直播间', {}).get('orders', 0), comp_stores.get('小米手表直播间', {}).get('gsv', 0)),
]
add_styled_table(doc,
    ['直播间', '订单数', 'GSV（元）'],
    [[r[0], f'{r[1]:,}', f'¥{r[2]:,.0f}'] for r in watch_rooms_comp] + [['合计', f'{sum(r[1] for r in watch_rooms_comp):,}', f'¥{sum(r[2] for r in watch_rooms_comp):,.0f}']],
    col_widths=[5, 3.5, 4]
)

doc.add_paragraph()
add_para(doc, '2.5 手表品类差距根因分析', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

reasons = [
    ('直播间数量差距：', '竞对手表号5个 vs 我方2个，直接导致流量入口不足。竞对单手表号平均4,804单，我方单手表号平均5,366单——我方单号效率更高，但号太少。'),
    ('千川投放差距：', '竞对千川投放消耗约为我方3倍，Watch 6是其重点投放品类。我方Watch 6千川ROI（18.9）高于手环（14.1），说明应加大Watch 6投放。'),
    ('商品卡运营不足：', 'Watch 6商品卡我方仅3,126单，竞对2,203单——虽然我方商品卡略多，但绝对量偏低。Watch 6用户有明确搜索需求，商品卡潜力大。'),
    ('高端手表（S系列）缺失：', '我方S系列仅366单（份额11.3%），竞对2,868单。高端手表的缺失意味着客单价和利润空间的流失。'),
    ('主播专业度差异：', '手表用户对手表功能、运动模式、健康监测等有更专业的需求，需要差异化的话术和主播能力。'),
]
for prefix, text in reasons:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_paragraph()

# ========================================
# 第三部分：全品类结构分析
# ========================================
add_title(doc, '三、全品类销售结构变化分析', level=1)

add_para(doc, '3.1 全品类销售结构', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

all_products = data['all_products']
sorted_products = sorted(all_products.items(), key=lambda x: x[1]['orders'], reverse=True)

add_styled_table(doc,
    ['品类', '订单数', 'GSV（元）', '占比', '结构定位'],
    [
        [name, f'{p["orders"]:,}', f'¥{p["gsv"]:,.0f}', f'{p["orders"]/data["total_orders"]*100:.1f}%',
         '绝对主力' if name == '小米手环10' else
         'GSV冠军' if name == 'REDMI Watch 6' else
         '利润担当' if name == '小米手环10 Pro' else
         '过渡产品' if name == '小米手环9 Pro' else
         '耳机引流' if 'Buds' in name else '']
        for name, p in sorted_products
    ],
    col_widths=[4, 2.5, 3.5, 2, 3]
)

doc.add_paragraph()
add_para(doc, '3.2 我方 vs 竞对品类结构差异', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_para(doc, '我方品类结构过于依赖手环10（占49.3%），而竞对结构更均衡。竞对在Watch 6（36.6%）和10 Pro（30.7%）上的占比远高于我方，说明竞对在高客单价产品的转化能力更强。')

our_total = our['orders']
comp_total = comp['orders']

add_styled_table(doc,
    ['品类', '我方订单', '我方占比', '竞对订单', '竞对占比', '差距分析'],
    [
        ['小米手环10', f'{our_products["小米手环10"]["orders"]:,}',
         f'{our_products["小米手环10"]["orders"]/our_total*100:.1f}%',
         f'{comp_products["小米手环10"]["orders"]:,}',
         f'{comp_products["小米手环10"]["orders"]/comp_total*100:.1f}%',
         '我方略优（+47.4%份额）'],
        ['REDMI Watch 6', f'{our_products["REDMI Watch 6"]["orders"]:,}',
         f'{our_products["REDMI Watch 6"]["orders"]/our_total*100:.1f}%',
         f'{comp_products["REDMI Watch 6"]["orders"]:,}',
         f'{comp_products["REDMI Watch 6"]["orders"]/comp_total*100:.1f}%',
         '严重落后（29.8%份额）'],
        ['小米手环10 Pro', f'{our_products["小米手环10 Pro"]["orders"]:,}',
         f'{our_products["小米手环10 Pro"]["orders"]/our_total*100:.1f}%',
         f'{comp_products["小米手环10 Pro"]["orders"]:,}',
         f'{comp_products["小米手环10 Pro"]["orders"]/comp_total*100:.1f}%',
         '严重落后（22.0%份额）'],
        ['小米手环9 Pro', f'{our_products["小米手环9 Pro"]["orders"]:,}',
         f'{our_products["小米手环9 Pro"]["orders"]/our_total*100:.1f}%',
         f'{comp_products["小米手环9 Pro"]["orders"]:,}',
         f'{comp_products["小米手环9 Pro"]["orders"]/comp_total*100:.1f}%',
         '落后（28.9%份额）'],
        ['Xiaomi 开放式耳机', f'{our_products["Xiaomi 开放式耳机"]["orders"]:,}',
         f'{our_products["Xiaomi 开放式耳机"]["orders"]/our_total*100:.1f}%',
         f'{comp_products["Xiaomi 开放式耳机"]["orders"]:,}',
         f'{comp_products["Xiaomi 开放式耳机"]["orders"]/comp_total*100:.1f}%',
         '持平（50.4%份额）'],
    ],
    col_widths=[3, 2.2, 2, 2.2, 2, 4]
)

doc.add_paragraph()

# ========================================
# 第四部分：竞对分析
# ========================================
add_title(doc, '四、竞对618销售表现及可取之处', level=1)

add_para(doc, '4.1 良米（主要竞对）核心策略', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

strategies_liangmi = [
    ('直播间矩阵优势：', '8个直播间形成密集矩阵，手环号（30,577单）和手表号（19,779单）是两个超级大号。单号平均产出8,500单，高于我方6,955单。'),
    ('千川"海投"策略：', '投放消耗约为我方3倍，以量取胜。虽然单视频CTR不如我方，但总播放量和成交规模更大。'),
    ('Watch 6重仓投入：', '良米Watch 6贡献31,440单、1,419.3万元GSV，是其最大品类，精准押注高GSV品类。'),
    ('10 Pro转化能力强：', '良米10 Pro贡献26,381单（我方仅7,458单），高客单价产品的转化话术值得学习。'),
    ('耳机独立运营：', '拥有专属耳机直播间（4,838单），耳机品类运营专业化。'),
    ('大促爆发力强：', '617→618增幅94%（我方86%），大促日的爆发节奏更好。'),
]
for prefix, text in strategies_liangmi:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_paragraph()
add_para(doc, '4.2 机械空间策略亮点', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

strategies_jixie = [
    ('极致效率运营：', '仅2个号就做到10,990单/号的产出，全场最高。'),
    ('"国补"概念营销：', '"小米智能穿戴国补号"主打国补优惠概念，精准切中价格敏感用户。'),
    ('纯直播模式：', '不做商品卡，集中资源打透直播渠道。'),
]
for prefix, text in strategies_jixie:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_paragraph()
add_para(doc, '4.3 竞对可取之处与融入方案', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_styled_table(doc,
    ['竞对做法', '可取之处', '融入我方运营的方案'],
    [
        ['良米：8号矩阵', '更多直播间=更多流量入口', '新增1-2个手表专属直播间'],
        ['良米：千川海投', '规模效应带来总成交优势', '保持CTR优势，逐步提升千川预算至竞对60-70%'],
        ['良米：Watch 6重仓', '精准押注高GSV品类', '将Watch 6作为下阶段重点品类'],
        ['良米：10 Pro强转化', '高客单价产品转化能力', '学习竞对10 Pro直播话术'],
        ['良米：耳机独立运营', '品类专业化提升效率', '整改违规后重新定位耳机号'],
        ['机械空间："国补"概念', '差异化概念精准获客', '直播间增加国补/以旧换新卖点'],
        ['竞对：大促爆发力', '617→618增幅94%', '建立大促预热SOP，提前7天蓄水'],
    ],
    col_widths=[4, 4.5, 6]
)

doc.add_paragraph()

# ========================================
# 第五部分：千川投放对比
# ========================================
add_title(doc, '五、千川视频投放对比分析', level=1)

add_para(doc, '5.1 核心指标对比（W1: 6/1-6/7）', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_styled_table(doc,
    ['指标', '我方', '竞对（良米）', '优劣势'],
    [
        ['有消耗视频数', '1,880条', '远超我方', '竞对规模更大'],
        ['总消耗', '¥58,238', '约¥175,000（3倍）', '竞对投入更大'],
        ['总成交金额', '¥876,254', '更高', ''],
        ['净成交订单', '2,249', '更高', ''],
        ['整体ROI', '15.05', '略有差异', '效率接近'],
        ['整体CTR', '7.23%', '低于我方', '★ 我方优势'],
        ['整体CVR', '1.55%', '低于我方', '★ 我方优势'],
        ['播放效率', '较低', '更高', '竞对规模优势'],
        ['AIGC投放', '大量使用', '几乎不用', '★ 我方差异化优势'],
    ],
    col_widths=[3, 3.5, 4, 4]
)

doc.add_paragraph()
add_para(doc, '5.2 关键结论', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

conclusions = [
    '我方CTR和CVR均高于竞对，说明内容质量和转化效率更优。',
    '竞对投放消耗约为我方3倍，以量取胜是其核心策略。',
    'AIGC动态创意是我方区别于竞对的差异化手段，竞对几乎未布局。',
    '竞对84.6%的视频ROI=0（无转化），存在大量低效投放，我方需引以为戒。',
    'Watch 6千川ROI（18.9）高于手环（14.1），应加大Watch 6投放力度。',
]
for c in conclusions:
    add_bullet(doc, c)

doc.add_paragraph()

# ========================================
# 第六部分：违规情况
# ========================================
add_title(doc, '六、违规情况及整改要求', level=1)

add_para(doc, '6.1 违规概况', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_para(doc, '全周期共21次违规，涉及5个直播间。耳机直播间违规最多（9次，占43%），是重点整改对象。')

add_styled_table(doc,
    ['直播间', '违规次数', '主要问题', '处罚情况'],
    [
        ['小米官方耳机直播间', '9次', '售后服务不符（5次）', '预警5次+违规3次+已撤销1次'],
        ['小米官方手表', '5次', '赠品信息不符（2次）', '预警5次'],
        ['小米官方手环直播间', '3次', '售后服务不符（2次）', '预警2次+已撤销1次'],
        ['小米手环10pro直播间', '3次', '售后服务不符（3次）', '预警3次'],
        ['小米数码旗舰店', '1次', '赠品信息不符（1次）', '预警1次'],
    ],
    col_widths=[4, 2, 4, 5]
)

doc.add_paragraph()
add_para(doc, '6.2 整改要求', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

rectifications = [
    '立即规范直播间话术，不得承诺"全国联保""终生售后""一年质保"等未确权表述。',
    '赠品信息需在商品详情页明确体现，直播间口播与详情页保持严格一致。',
    '建立直播话术审核机制，对每场直播的宣传文案进行提前审核。',
    '直播间任何抽奖/赠品活动必须使用平台官方福袋工具，严禁第三方抽奖。',
    '对小米官方耳机直播间进行专项整改，该直播间已出现账号级处罚。',
]
for r in rectifications:
    add_bullet(doc, r)

doc.add_paragraph()

# ========================================
# 第七部分：对策与行动计划
# ========================================
add_title(doc, '七、手表品类专项对策与行动计划', level=1)

add_para(doc, '7.1 手表品类对策', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_styled_table(doc,
    ['序号', '对策', '优先级', '预期效果', '负责方'],
    [
        ['1', '增加手表专属直播间数量（参考良米5个号布局）', '高', '+5,000-8,000单/月', '运营团队'],
        ['2', '加强Watch 6千川视频投放（ROI 18.9高于手环14.1）', '高', '放大ROI优势', '投放团队'],
        ['3', 'Watch 6商品卡专项优化（标题/主图/详情页）', '中', '+2,000-3,000单', '电商运营'],
        ['4', '手表S系列直播话术升级（强调工艺/材质/场景）', '中', '提升高端转化', '主播团队'],
        ['5', '手表品类专属主播培养', '中', '提升转化率', '主播团队'],
    ],
    col_widths=[1.2, 7, 2, 3.5, 2.5]
)

doc.add_paragraph()
add_para(doc, '7.2 耳机品类对策', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_styled_table(doc,
    ['序号', '对策', '优先级', '预期效果'],
    [
        ['1', '立即整改耳机直播间违规问题', '紧急', '恢复账号健康度'],
        ['2', '精简耳机SKU，聚焦开放式耳机+Buds 8青春版+Buds 8 Pro', '高', '提升单品转化效率'],
        ['3', '在手环/手表直播间增加耳机搭售推荐', '高', '+500-1,000单/月'],
        ['4', '开放式耳机作为差异化重点（份额50.4%）', '中', '巩固优势品类'],
    ],
    col_widths=[1.2, 8.5, 2, 3.5]
)

doc.add_paragraph()
add_para(doc, '7.3 行动计划时间表', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_styled_table(doc,
    ['时间', '行动项', '目标'],
    [
        ['1个月内', '耳机直播间违规整改完成', '消除账号风险'],
        ['1个月内', '建立标准化话术库（含合规审核）', '杜绝违规'],
        ['1个月内', 'Watch 6千川投放预算提升50%', '放大手表ROI优势'],
        ['1个月内', '10 Pro直播间话术优化（参考竞对）', '提升高客单价转化'],
        ['1-3个月', '新增1-2个手表专属直播间', '补齐直播间矩阵'],
        ['1-3个月', '商品卡运营体系搭建', '商品卡占比提升至25%+'],
        ['1-3个月', '千川投放预算提升至竞对60%水平', '缩小投放规模差距'],
        ['3-6个月', '直播间矩阵完善至8-10个', '全面对标竞对'],
        ['3-6个月', '手表品类份额提升至40%+', '缩小手表差距'],
    ],
    col_widths=[3, 7, 5]
)

doc.add_paragraph()
add_para(doc, '7.4 后续尝试与协作需求', bold=True, size=12, color=RGBColor(0x1E, 0x90, 0xFF))

add_styled_table(doc,
    ['方向', '具体内容', '需要的协作支持'],
    [
        ['直播间矩阵扩展', '新增手表专属直播间1-2个', '主播资源调配、直播间资质申请'],
        ['千川规模化', '投放预算提升至竞对60-70%', '投放预算审批、素材批量生产'],
        ['AIGC内容深化', '扩大AIGC动态创意视频比例', 'AIGC工具授权、创意脚本储备'],
        ['商品卡精细化', '建立商品卡专属运营团队', '运营人员配置、详情页设计资源'],
        ['大促爆发力', '建立大促前7天预热SOP', '大促专项预算、限时优惠机制'],
        ['合规话术体系', '建立标准化话术库', '法务/合规审核支持'],
        ['竞对实时监控', '建立竞对直播间数据监控', '数据抓取工具、第三方平台'],
        ['用户画像分析', '分析手环/手表/耳机用户差异', '抖音数据银行/巨量云图权限'],
    ],
    col_widths=[3.5, 5.5, 6]
)

# ===== 结尾 =====
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('— 文档完 —')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('文档生成时间：2026年6月28日')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ========== 保存 ==========
output_path = r'C:\Users\Administrator\Desktop\小米手环直播间销量分析\618复盘分析总结文档.docx'
doc.save(output_path)
print(f"Word文档已生成：{output_path}")
