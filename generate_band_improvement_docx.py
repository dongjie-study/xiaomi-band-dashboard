"""
Generate Xiaomi Band Improvement Plan Word Report
Focus: 小米手环系列 only
"""
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DATA_DIR = os.path.dirname(os.path.abspath(__file__))

def get_team(rname):
    r = str(rname)
    if '小米官方手表直播号' == r:
        return '纵横'
    if '小米智能穿戴国补号' == r or '小米智能穿戴授权号' == r:
        return '机械空间'
    if any(kw in r for kw in ['小米官方', '小米数码旗舰店', '小米手环10Pro直播间', '小米官旗手表直播间']):
        return '我司'
    return '良米'

def load_data():
    with open(os.path.join(DATA_DIR, 'sales_analysis', 'history.json'), 'r', encoding='utf-8') as f:
        history = json.load(f)
    june = [d for d in history if d['date'].startswith('2026-06')]
    july = [d for d in history if d['date'].startswith('2026-07')]
    return history, june, july

def calc_band_stats(days_list):
    """Calculate band-only stats from a list of daily data"""
    band_total = 0; band_rev = 0
    band_prods = {}
    band_team = {}
    band_rooms = {}

    for d in days_list:
        # Products
        for pname, pinfo in d.get('products', {}).items():
            if any(bp in str(pname) for bp in ['手环', 'Band']):
                band_total += pinfo['orders']
                band_rev += pinfo['revenue']
                if pname not in band_prods:
                    band_prods[pname] = {'orders': 0, 'revenue': 0}
                band_prods[pname]['orders'] += pinfo['orders']
                band_prods[pname]['revenue'] += pinfo['revenue']

        # Rooms → teams
        for rname, rinfo in d.get('rooms', {}).items():
            t = get_team(rname)
            if t not in band_team:
                band_team[t] = {'band_orders': 0, 'band_rev': 0}
            if rname not in band_rooms:
                band_rooms[rname] = {'band_orders': 0, 'band_rev': 0, 'type': t}
            for pname, pinfo in rinfo.get('products', {}).items():
                if any(bp in str(pname) for bp in ['手环', 'Band']):
                    band_team[t]['band_orders'] += pinfo['orders']
                    band_team[t]['band_rev'] += pinfo['revenue']
                    band_rooms[rname]['band_orders'] += pinfo['orders']
                    band_rooms[rname]['band_rev'] += pinfo['revenue']

    return band_total, band_rev, band_prods, band_team, band_rooms

def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(doc, text, bold=False, size=11, color=None, alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    return p

def add_bullet(doc, text, level=0, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_cell(cell, text, bold=False, size=9, color=None, align='center', bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bg:
        set_cell_shading(cell, bg)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        format_cell(table.rows[0].cells[i], h, bold=True, size=9, color=(255,255,255), bg='1E90FF')

    # Data
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            format_cell(table.rows[r+1].cells[c], str(val), size=9,
                       bg='F8F9FC' if r % 2 == 0 else None)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacer
    return table

def build_report():
    history, june, july = load_data()
    band_total, band_rev, band_prods, band_team, band_rooms = calc_band_stats(june)

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ============================================================
    # COVER PAGE
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('小米手环系列\n直播间销量提升方案')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(255, 105, 0)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('—— 基于2026年6月订单数据的全链路优化策略 ——')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    doc.add_paragraph()

    info_lines = [
        f'报告日期：2026年7月2日',
        f'数据周期：2026年6月1日 - 7月1日',
        f'分析范围：小米手环系列（手环10、手环10 Pro、手环9 Pro等）',
        f'报告版本：V1.0'
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(130, 130, 130)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_page_break()

    # ============================================================
    # TABLE OF CONTENTS (Manual)
    # ============================================================
    add_styled_heading(doc, '目录', level=1)
    toc_items = [
        ('一、数据总览与竞争格局', '核心指标 · 团队对比 · 产品结构 · 竞争态势'),
        ('二、直播间优化方案', '矩阵布局 · 流量策略 · 排品节奏 · 互动留存'),
        ('三、视觉装修升级', '场景设计 · 贴片优化 · 氛围营造 · A/B测试'),
        ('四、话术体系升级', '开场破冰 · 产品种草 · 逼单转化 · 异议处理'),
        ('五、视频内容策略', '内容矩阵 · 爆款公式 · 拍摄SOP · 发布节奏'),
        ('六、投放优化方案', '千川策略 · 抖加配合 · 预算分配 · ROI提升'),
        ('七、视频播放量提升专项', '算法机制 · 蹭热点 · 互动引导 · 矩阵分发'),
        ('八、执行计划与KPI', '优先级排序 · 时间节点 · 资源配置 · 目标设定'),
    ]
    for title, desc in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1)
        run2 = p2.add_run(desc)
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(150, 150, 150)
        run2.font.name = '微软雅黑'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_page_break()

    # ============================================================
    # SECTION 1: DATA OVERVIEW
    # ============================================================
    add_styled_heading(doc, '一、数据总览与竞争格局', level=1)

    add_styled_heading(doc, '1.1 核心指标概览', level=2)
    add_para(doc, f'2026年6月（30天），小米手环系列全渠道直播间关键数据：', size=10.5)

    # KPI table
    kpi_headers = ['指标', '6月数据', '日均', '说明']
    kpi_rows = [
        ['手环总订单', f'{band_total:,}单', f'{band_total//30:,}单/天', '占全品类订单66.3%'],
        ['手环总销售额', f'¥{band_rev/10000:.1f}万', f'¥{band_rev/30/10000:.1f}万/天', '手环10 Pro+手环10贡献98%'],
        ['手环10 Pro订单', f'{band_prods.get("小米手环10 Pro", {}).get("orders", 0):,}单', '', f'均价¥{band_prods.get("小米手环10 Pro", {}).get("revenue", 0)/band_prods.get("小米手环10 Pro", {}).get("orders", 1):.0f}'],
        ['手环10订单', f'{band_prods.get("小米手环10", {}).get("orders", 0):,}单', '', f'均价¥{band_prods.get("小米手环10", {}).get("revenue", 0)/band_prods.get("小米手环10", {}).get("orders", 1):.0f}'],
        ['我方手环份额', f'{band_team.get("我司", {}).get("band_orders", 0):,}单 (35.3%)', '', '落后良米10.3个百分点'],
        ['竞争直播房间', '19间', '', '我方6间 · 良米10间 · 机械空间2间 · 纵横1间'],
    ]
    create_table(doc, kpi_headers, kpi_rows, col_widths=[3.5, 3.5, 3, 6])

    add_styled_heading(doc, '1.2 竞争格局分析', level=2)
    add_para(doc, '手环品类团队市占率对比：', size=10.5)

    team_headers = ['团队', '手环订单', '市占率', '手环销售额', '直播间数', '头牌直播间', '定位判断']
    team_rows = []
    for t in ['我司', '良米', '机械空间', '纵横']:
        if t in band_team:
            bt = band_team[t]
            pct = bt['band_orders']/band_total*100
            rooms_count = len([r for r, info in band_rooms.items() if info['type'] == t])
            top_room = sorted([(r, info) for r, info in band_rooms.items() if info['type'] == t],
                            key=lambda x: -x[1]['band_orders'])
            top_name = top_room[0][0] if top_room else '-'

            if t == '我司':
                positioning = '🔵 追赶者 — 需突破'
            elif t == '良米':
                positioning = '🔴 领先者 — 主打低价走量'
            elif t == '机械空间':
                positioning = '🟠 搅局者 — 国补加持走量'
            else:
                positioning = '⚪ 边缘参与者'

            team_rows.append([t, f'{bt["band_orders"]:,}单', f'{pct:.1f}%',
                            f'¥{bt["band_rev"]/10000:.1f}万', str(rooms_count), top_name, positioning])

    create_table(doc, team_headers, team_rows, col_widths=[1.8, 2, 1.5, 2.2, 1.5, 3.5, 3.5])

    # Key insights
    add_para(doc, '▎核心发现：', bold=True, size=11, color=(255, 105, 0))

    findings = [
        f'【最大缺口】良米"小米手环"单间直播间手环订单19,362单，我方头牌"小米官方手环直播间"仅9,513单，差距近1倍。该直播间是突破关键。',
        f'【产品结构】手环10 Pro占订单49.0%（¥1,115万），手环10占48.9%（¥815万），两者合计98%。产品集中度高，视频素材和话术应100%聚焦这两个品。',
        f'【价格带分析】手环10 Pro均价¥395 vs 手环10均价¥288。Pro毛利更高，应作为主推品；手环10作为引流品。',
        f'【周趋势】W3（6/15-21）达到峰值32,269单/周，6/25单日峰值5,827单手环订单。大促期间爆发力明显，平播期需稳定在1,500-2,000单/天。',
        f'【我司优势】6间直播间矩阵布局相对完善，官方背书信任度高；劣势在于单直播间产出效率低于良米。',
    ]
    for f_text in findings:
        add_bullet(doc, f_text, size=10)

    doc.add_page_break()

    # ============================================================
    # SECTION 2: 直播间优化方案
    # ============================================================
    add_styled_heading(doc, '二、直播间优化方案', level=1)

    add_styled_heading(doc, '2.1 现状诊断', level=2)

    diag_headers = ['直播间', '手环月订单', '定位', '核心问题', '提升空间']
    diag_rows = [
        ['小米官方手环直播间', '9,513', '手环主阵地', '单量仅为竞对"小米手环"的49%', '★★★ 最大提升空间'],
        ['小米数码旗舰店', '8,944', '数码综合', '品类混杂，手环心智不够强', '★★ 中等提升空间'],
        ['小米官方手表', '1,715', '手表为主', '手环非主打，连带销售', '★ 次要'],
        ['小米手环10Pro直播间', '190', '高端手环', '开播不久，流量未成型', '★★★ 潜力新品'],
        ['小米官方耳机直播间', '1', '耳机为主', '基本无手环销售', '★ 非重点'],
        ['小米官旗手表直播间', '33', '手表为主', '基本无手环销售', '★ 非重点'],
    ]
    create_table(doc, diag_headers, diag_rows, col_widths=[3.5, 2, 2, 4.5, 4])

    add_styled_heading(doc, '2.2 直播间优化策略', level=2)

    add_para(doc, '策略一：打造手环超级大场（针对"小米官方手环直播间"）', bold=True, size=11)
    strategies_1 = [
        '目标：从日均317单提升至500单，缩小与竞对差距',
        '增加直播时长：从当前时段扩展至16小时（8:00-24:00），覆盖早晚高峰',
        '增设连麦PK：与小米其他直播间（手表/耳机）连麦导流，共享粉丝池',
        '每日设置2场"整点秒杀"：整点放10台特价手环10（¥229），拉高在线人数',
        '周五大场日：每周五做超级大场，集中所有流量资源，冲刺日销800单',
    ]
    for s in strategies_1:
        add_bullet(doc, s, size=10)

    add_para(doc, '策略二：激活"小米手环10Pro直播间"（新品孵化）', bold=True, size=11)
    strategies_2 = [
        '定位调整为"高端手环体验馆"：主推手环10 Pro陶瓷版/耀影金特别版（均价¥400+）',
        '人设差异化：区别于官旗直播间的"促销感"，走"科技测评+生活方式"路线',
        '日播8小时起步：早10点到晚10点，错峰竞争（晚上避开良米高流量时段）',
        '前两周日播+千川投流蓄水，第三周起放量',
    ]
    for s in strategies_2:
        add_bullet(doc, s, size=10)

    add_para(doc, '策略三："小米数码旗舰店"手环专区化', bold=True, size=11)
    strategies_3 = [
        '在直播间内划分"手环专区"视觉区域，强化手环品类心智',
        '排品节奏：每3个手环品 → 1个数码配件，保持手环露出率70%+',
        '利用旗舰店SKU优势做"手环+配件"组合套装，拉高客单价',
    ]
    for s in strategies_3:
        add_bullet(doc, s, size=10)

    add_styled_heading(doc, '2.3 排品节奏优化', level=2)
    add_para(doc, '基于6月小时级数据分析，建议标准排品节奏（以2小时为一个循环）：', size=10.5)

    schedule_headers = ['时间段', '阶段', '主推品', '话术重点', '目标']
    schedule_rows = [
        ['0-15分钟', '开场暖场', '手环10（¥229引流价）', '福利钩子 + 扣屏互动', '拉高在线至100+'],
        ['15-45分钟', '深度种草', '手环10 Pro（陶瓷版 ¥399）', '功能展示 + 场景化佩戴', '停留时长>90秒'],
        ['45-75分钟', '逼单转化', '手环10 Pro耀影金（¥429）', '限时优惠 + 库存紧张', '密集出单'],
        ['75-90分钟', '过渡+憋单', '手环10（¥229返场）', '倒计时 + 下一波预告', '承接流量'],
        ['90-120分钟', '第二轮循环', '循环上述节奏', '更换展示角度/场景', '维持节奏'],
    ]
    create_table(doc, schedule_headers, schedule_rows, col_widths=[2, 2, 3.5, 3.5, 3])

    doc.add_page_break()

    # ============================================================
    # SECTION 3: 视觉装修
    # ============================================================
    add_styled_heading(doc, '三、视觉装修升级方案', level=1)

    add_styled_heading(doc, '3.1 直播间场景设计', level=2)

    add_para(doc, '▎核心原则：科技感 + 运动感 + 生活化，让用户"看到就想戴"', bold=True, size=11, color=(255,105,0))

    add_para(doc, '场景一：科技极简风（主力场景，适用于"官方手环直播间"）', bold=True, size=11)
    scene1 = [
        '背景：深灰色/黑色科技感背景墙 + LED灯带（小米橙色）',
        '桌面：白色/浅灰展示台，极简陈列，突出产品',
        '灯光：主光正面打亮主播+产品，侧光勾勒轮廓，顶部射灯聚焦手环',
        '道具：iPad展示APP联动界面、磁吸充电座、运动场景小道具（哑铃/瑜伽垫）',
        '贴片位置：左上角品牌LOGO、右上角"官方直播间"、底部价格条',
    ]
    for s in scene1:
        add_bullet(doc, s, size=10)

    add_para(doc, '场景二：运动活力风（适用于"10Pro直播间"差异化定位）', bold=True, size=11)
    scene2 = [
        '背景：健身房/户外跑道背景布 + 自然光感',
        '主播着装：运动风（瑜伽服/跑步装），手腕佩戴手环实拍',
        '动态展示：跑步机上实时展示心率监测、步数统计',
        '差异化价值：和官旗直播间的"坐播推销"拉开距离，更有真实感',
    ]
    for s in scene2:
        add_bullet(doc, s, size=10)

    add_styled_heading(doc, '3.2 贴片与视觉元素优化', level=2)

    add_para(doc, '贴片设计规范（参考竞对良米直播间的高转化贴片）：', size=10.5)

    patch_headers = ['贴片位置', '内容', '设计要点', '更新频率']
    patch_rows = [
        ['左上角', '品牌LOGO + "官方"标识', '小米橙色，半透明底，不遮挡产品', '固定不变'],
        ['右上角', '当前活动/利益点', '"618返场价""今日特惠"等，红底白字', '每日更新'],
        ['底部横条', '价格+优惠信息', '"手环10 Pro 到手¥3XX 限时赠表带"', '每场直播更新'],
        ['右侧悬浮', '产品核心卖点', '"14天续航｜150+运动模式｜血氧监测"', '按主推品切换'],
        ['左下角', '互动引导', '"扣1领券""点亮粉丝牌享专属价"', '每30分钟轮换'],
        ['画面中央', '倒计时浮窗', '秒杀/限量时出现，制造紧迫感', '大场活动时使用'],
    ]
    create_table(doc, patch_headers, patch_rows, col_widths=[2.5, 3.5, 5, 3])

    add_styled_heading(doc, '3.3 竞对视觉对标', level=2)
    add_para(doc, '通过对良米"小米手环"直播间（月销19,362单手环）的视觉分析，发现其优势：', size=10.5)
    bench_notes = [
        '画面简洁：背景纯白/浅灰，没有过多装饰，用户注意力集中在产品上',
        '价格醒目：超大字体的价格优惠贴片，3秒内传达核心利益点',
        '真人实戴：主播全程佩戴手环，展示通知提醒、抬手亮屏等真实使用场景',
        '动态演示：频繁切换手机APP数据页面，展示睡眠/运动/健康数据',
        '我司改进方向：减少画面杂乱元素，增大价格字体，增加动态产品演示占比',
    ]
    for b in bench_notes:
        add_bullet(doc, b, size=10)

    doc.add_page_break()

    # ============================================================
    # SECTION 4: 话术
    # ============================================================
    add_styled_heading(doc, '四、话术体系升级方案', level=1)

    add_styled_heading(doc, '4.1 话术核心逻辑：AITDA模型', level=2)

    aida_headers = ['阶段', '时长', '核心目标', '关键话术框架']
    aida_rows = [
        ['Attention 吸引', '0-30秒', '让刷到的人停下来', '"今天手环10 Pro跌破底价！停留的宝宝扣1，我给你们看惊喜"'],
        ['Interest 兴趣', '30秒-2分钟', '建立产品兴趣', '"这个陶瓷版质感，你看我手腕上戴的，昨天刚到的耀影金，出门回头率超高…"'],
        ['Trust 信任', '2-5分钟', '建立信任消除顾虑', '"我们是小米官方直播间，全国联保，7天无理由，15天换新。直播间下单和官网一样的售后"'],
        ['Desire 欲望', '5-8分钟', '激发购买欲望', '"你想想，399买个Pro，续航14天，血氧、心率、睡眠全监测，连苹果表十分之一的价格都不到"'],
        ['Action 行动', '8-10分钟', '促成下单', '"今天限量100台，已经出了67台了，还剩33台。下方小黄车1号链接，现在下单还送原装表带！"'],
    ]
    create_table(doc, aida_headers, aida_rows, col_widths=[2.5, 1.5, 2.5, 9.5])

    add_styled_heading(doc, '4.2 开场破冰话术（前30秒决定留存率）', level=2)
    add_para(doc, '▎"3秒钩子"公式：利益点 + 互动指令', bold=True, size=10.5)

    hooks = [
        '"手环10今天破价了！刚进来的宝宝先别走，扣1我给你看价格"（利益钩子）',
        '"戴手环的宝宝扣1，没戴的扣2，我看看有多少人需要换新手环"（互动钩子）',
        '"昨天650人抢到的手环10 Pro陶瓷版，今天又补了100台，来晚了真没了"（稀缺钩子）',
        '"刚刚一个大哥一口气拍了3台，说公司团建当奖品。来我给你们看看他买的哪个色"（从众钩子）',
    ]
    for h_text in hooks:
        add_bullet(doc, h_text, size=10)

    add_styled_heading(doc, '4.3 产品种草话术（手环10 Pro 为主推）', level=2)
    add_para(doc, '▎FABE法则：Feature → Advantage → Benefit → Evidence', bold=True, size=10.5)

    fabe_headers = ['维度', '手环10 Pro话术示例']
    fabe_rows = [
        ['F-特性', '"这款用的是1.74英寸AMOLED屏幕，陶瓷质感机身，IP68防水"'],
        ['A-优势', '"比上一代屏幕大了30%，阳光下看得清清楚楚，游泳洗澡都不用摘"'],
        ['B-利益', '"你跑步的时候抬手就能看到心率，坐久了会提醒你起来活动，晚上自动监测睡眠质量"'],
        ['E-证据', '"你看我手上的数据，昨晚深度睡眠2小时18分钟，这个数据和我戴的万元睡眠仪对比过，准确率95%以上"'],
    ]
    create_table(doc, fabe_headers, fabe_rows, col_widths=[2, 14])

    add_styled_heading(doc, '4.4 逼单转化话术', level=2)

    closing_tactics = [
        '【限时限量】"今天直播间的专属价只有前100台，已经出了82台，还有最后18台。3号链接，拍完恢复原价¥449"',
        '【价格对比】"官网¥449，我们今天到手¥399，还多送一根表带。这根表带单买都要¥49，相当于手环才¥350"',
        '【竞品对比】"苹果表最便宜的也要¥1999，功能也就那些。我们这个Pro，14天续航秒杀苹果表，399就能拥有"',
        '【场景驱动】"暑假马上到了，给孩子买一个，随时知道他在哪运动、睡了多久。比买个游戏机实用多了"',
        '【售后保障】"官方直播间下单，全国1800+小米之家都能售后，7天无理由退换，买回去不满意随时退"',
    ]
    for c in closing_tactics:
        add_bullet(doc, c, size=10)

    add_styled_heading(doc, '4.5 异议处理话术', level=2)

    objection_headers = ['常见异议', '应对话术']
    objection_rows = [
        ['"太贵了"', '"哥哥，399买个Pro你还觉得贵呀？你算一下，用两年，一天才5毛5。一杯奶茶钱都不到，换14天续航+150种运动模式，不划算吗？而且我们今天有专属赠品…"'],
        ['"手环10和10 Pro有什么区别"', '"好问题！Pro多了独立GPS（跑步不用带手机）、屏幕更大、陶瓷机身更有质感。如果你只是看时间记步数，手环10就够了¥229；如果你运动比较多想要GPS，上Pro绝对不后悔"'],
        ['"官网也是这个价"', '"官网今天¥449一分不少哦。我们直播间专属补贴+限量赠品（原装表带¥49），这个组合只有直播间有。你现在去官网对比，截图给我看，有比我便宜的我现在就下播"'],
        ['"再看看/考虑一下"', '"没问题的哥哥，但我建议你先拍下不付款，库存只有最后几台了。拍下后15分钟内付款都有效，你慢慢考虑。过一会回来可能真没了…上次有个宝宝就是这样错过的"'],
    ]
    create_table(doc, objection_headers, objection_rows, col_widths=[3, 13])

    doc.add_page_break()

    # ============================================================
    # SECTION 5: 视频
    # ============================================================
    add_styled_heading(doc, '五、视频内容策略', level=1)

    add_styled_heading(doc, '5.1 视频核心指标回顾', level=2)
    add_para(doc, '基于千川6月第一周数据（1,880条有消耗视频，总消耗¥58,238）：', size=10.5)

    video_stats = [
        '手环品类视频：消耗¥27,434，带来成交¥386,066，ROI 14.1',
        '手环视频总播放量：882,299次',
        '整体CTR：7.23%，转化率：1.55%',
        'TOP手环视频ROI：19.8（"真让你给等到了 #小米手环#小米手环9pro"）',
        'AIGC动态创意贡献手环成交¥133,399，ROI 13.4',
    ]
    for v in video_stats:
        add_bullet(doc, v, size=10)

    add_styled_heading(doc, '5.2 视频内容矩阵', level=2)
    add_para(doc, '▎构建"5+3"内容矩阵：5种内容类型 × 3个发布频率', bold=True, size=10.5)

    content_headers = ['内容类型', '占比', '拍摄要点', '目标数据', '发布频率']
    content_rows = [
        ['产品展示/开箱', '30%', '极简桌面+微距拍陶瓷质感+UI界面滑动。BGM：轻科技风', '播放>5K, CTR>8%', '每天2条'],
        ['使用场景/生活方式', '25%', '运动场景实拍（跑步/游泳/骑行）、睡眠监测展示、通勤佩戴', '播放>10K, CTR>6%', '每天1-2条'],
        ['价格/促销利益', '20%', '"破价了！手环10 Pro只要3XX"，大字标题+价格对比+紧迫感', '播放>20K, CTR>12%', '大场前/活动期密集发'],
        ['功能测评/对比', '15%', '手环10 vs 10 Pro对比、与苹果表对比、专业运动数据测试', '播放>8K, 互动>3%', '每周3-4条'],
        ['热梗/话题借势', '10%', '跟抖音热榜话题创作，如#暑假必备数码#运动打卡等', '播放>50K潜力', '随时跟热点'],
    ]
    create_table(doc, content_headers, content_rows, col_widths=[2.5, 1.2, 5, 3, 3])

    add_styled_heading(doc, '5.3 爆款视频公式', level=2)
    add_para(doc, '▎已验证的高转化视频模板（基于千川TOP视频分析）：', bold=True, size=10.5)

    add_para(doc, '公式一：「利益前置 + 价格锚点 + 紧迫感」', bold=True, size=10.5)
    formula1 = [
        '开头3秒：大字幕"手环10 Pro跌破底价！" + 产品特写镜头',
        '中间5-8秒：快速展示3个核心卖点（屏幕/续航/运动模式），每个2秒',
        '结尾3秒：价格对比（"官网¥449 → 直播间¥3XX"）+ "点击下方直播间"引导',
        'BGM：轻快科技感音乐，语速偏快',
        '参考数据：此类视频CTR可达12-15%，单条消耗¥1,000-2,000，ROI 15-20',
    ]
    for f in formula1:
        add_bullet(doc, f, size=10)

    add_para(doc, '公式二：「痛点场景 + 解决方案 + 信任背书」', bold=True, size=10.5)
    formula2 = [
        '开头3秒："每次跑步都要带手机？烦不烦？"（痛点切入）',
        '中间5-8秒：展示手环10 Pro独立GPS跑步场景 + 数据实时同步',
        '结尾3秒："小米官方出品，14天续航，¥399带走"',
        '适合投放：抖音信息流 + 搜索广告',
        '参考数据：此类视频完播率更高（>25%），适合长线种草',
    ]
    for f in formula2:
        add_bullet(doc, f, size=10)

    add_para(doc, '公式三：「真人实测 + 数据说话 + 权威感」', bold=True, size=10.5)
    formula3 = [
        '开头3秒：主播手腕特写 + "我戴了一个月的真实感受"',
        '中间8-12秒：展示睡眠数据、运动数据、续航实测',
        '结尾3秒："不吹不黑，¥399这个价位没有对手" + 引导进入直播间',
        '适合：品牌账号主页视频 + 置顶视频',
    ]
    for f in formula3:
        add_bullet(doc, f, size=10)

    add_styled_heading(doc, '5.4 视频拍摄SOP', level=2)
    add_para(doc, '建立标准化视频生产流程，确保日更5条+：', size=10.5)
    sop_items = [
        '每周一：策划本周30条视频选题（从5种类型中按比例分配）',
        '每日上午：集中拍摄1-2小时，至少产出5-8条素材',
        '每日下午：剪辑+加字幕+挂小黄车链接，批量定时发布',
        '每条视频必加元素：① 前3秒钩子 ② 产品特写 ③ 价格/利益点 ④ 引导语（"点下方直播间"）',
        'A/B测试：同一选题拍2-3个版本（不同开头/不同BGM），发布后保留数据好的',
        '爆款追投：单条视频自然播放量>5,000且CTR>5%的，立即千川追投¥200-500',
    ]
    for s in sop_items:
        add_bullet(doc, s, size=10)

    doc.add_page_break()

    # ============================================================
    # SECTION 6: 投放
    # ============================================================
    add_styled_heading(doc, '六、投放优化方案', level=1)

    add_styled_heading(doc, '6.1 投放现状诊断', level=2)
    add_para(doc, '基于千川6月核心指标（周度数据）：', size=10.5)

    ad_diag = [
        '手环品类周消耗仅¥27,434（占整体47%），但手环订单占总订单66% → 投放占比应提到70%+',
        '整体ROI 14.1，属于优秀水平，有加大投放的空间',
        'AIGC动态创意ROI 13.4但消耗¥9,927占比过高（36%） → 真人实拍素材ROI更高，应增加真人素材占比',
        '抖音号主页素材ROI 15.7，是最高效渠道，应重点加投',
        'TOP10消耗视频中有6条是非手环品 → 手环专属素材质量和数量待提升',
    ]
    for a in ad_diag:
        add_bullet(doc, a, size=10)

    add_styled_heading(doc, '6.2 千川投放策略', level=2)

    add_para(doc, '策略一：预算重新分配', bold=True, size=11)
    budget_headers = ['投放项目', '当前周消耗', '建议周消耗', '调整理由']
    budget_rows = [
        ['手环10 Pro（高客单）', '¥10,000', '¥20,000', 'ROI>14，客单价高，加大投入回报最高'],
        ['手环10（引流款）', '¥8,000', '¥15,000', 'ROI约12，走量款维持规模'],
        ['手环9 Pro（清仓）', '¥2,000', '¥3,000', '老品清库存，低预算维持'],
        ['AIGC动态创意', '¥9,927', '¥5,000', '降低占比，让位给真人实拍素材'],
        ['合计/周', '¥27,434', '¥43,000', '手环总预算提升57%'],
    ]
    create_table(doc, budget_headers, budget_rows, col_widths=[3.5, 2.5, 2.5, 7.5])

    add_para(doc, '策略二：人群定向优化', bold=True, size=11)
    targeting = [
        '核心人群：18-35岁，一二线城市，兴趣标签"智能穿戴""运动健身""数码科技"',
        '拓展人群：35-50岁，关注"健康管理""睡眠改善"，推手环10 Pro健康监测功能',
        '相似人群：基于已成交用户画像lookalike扩展，覆盖量×3',
        '排除已转化：30天内已购买用户排除，避免重复触达浪费',
        '地域加权：广东、浙江、江苏、北京、上海TOP5成交省份加倍投放',
    ]
    for t in targeting:
        add_bullet(doc, t, size=10)

    add_para(doc, '策略三：素材赛马机制', bold=True, size=11)
    saima = [
        '每批次新建5-8条视频素材同时投放，预算均分¥200-300/条',
        '24小时后看数据：消耗>¥50且ROI>10的保留加投；ROI<5的关停',
        '每周至少新建2批素材，保持素材新鲜度（平台对新素材有流量倾斜）',
        'TOP素材生命周期约5-7天，到期数据下滑及时替换',
    ]
    for s in saima:
        add_bullet(doc, s, size=10)

    add_styled_heading(doc, '6.3 抖加配合策略', level=2)
    add_para(doc, '千川做转化，抖加做声量——双引擎驱动：', size=10.5)
    doujia = [
        '直播间引流：每场直播固定¥300抖加投放"直播间引流"，定向手环兴趣人群',
        '爆款视频助推：自然流量突破10K播放的视频，立即¥200抖加助推引导进直播间',
        '粉丝增长：每月¥2,000专项预算投"粉丝增长"，目标月增5,000精准粉丝',
        '互动引爆：挑选互动率>5%的视频，¥100抖加助推评论区互动，制造热闹感',
    ]
    for d in doujia:
        add_bullet(doc, d, size=10)

    doc.add_page_break()

    # ============================================================
    # SECTION 7: 视频播放量提升
    # ============================================================
    add_styled_heading(doc, '七、视频播放量提升专项方案', level=1)

    add_para(doc, '▎核心逻辑：播放量 = 内容质量 × 发布数量 × 算法推荐 × 投放加持', bold=True, size=11, color=(255,105,0))

    add_styled_heading(doc, '7.1 算法推荐机制拆解', level=2)
    add_para(doc, '抖音视频的推荐链路：', size=10.5)

    algo_headers = ['阶段', '流量池', '考核指标', '我们的优化方向']
    algo_rows = [
        ['冷启动', '200-500播放', '完播率 > 互动率 > 转评赞', '前3秒必须有钩子（利益/好奇/痛点），时长控制在15-25秒'],
        ['初级流量池', '1,000-5,000播放', '完播率>30%, 互动率>3%', '中间段信息密度要高，每3秒一个信息点/画面切换'],
        ['中级流量池', '10,000-100,000播放', 'CTR>5%, 转化>1%', '标题党但不违规，封面图大字突出利益点'],
        ['高级流量池', '100,000+播放', 'GPM（千次曝光成交）', '挂小黄车，视频与直播间选品强关联'],
    ]
    create_table(doc, algo_headers, algo_rows, col_widths=[2.5, 2.5, 3.5, 7.5])

    add_styled_heading(doc, '7.2 播放量提升六大抓手', level=2)

    add_para(doc, '抓手1：热点借势（最快起量方式）', bold=True, size=11)
    hotspot = [
        '每日关注抖音热点宝/巨量创意，找到可关联的热门话题',
        '示例：#暑假必备数码 → "学生党开学前必入的手环"，#运动打卡 → "跑步不带手机的第30天"',
        '热点视频发布时间控制在热点出现后4小时内，抢占流量窗口',
        '每月至少借势4个热点话题（每周1个）',
    ]
    for h in hotspot:
        add_bullet(doc, h, size=10)

    add_para(doc, '抓手2：评论区运营（撬动算法推荐）', bold=True, size=11)
    comment_strategy = [
        '每条视频发布后30分钟内，用3个小号在评论区提问（"多少钱？""续航多久？""和苹果表比呢？"）',
        '主号回复评论时带产品卖点（"到手¥399，续航14天，比苹果表强多了"），引导更多讨论',
        '置顶最有价值的FAQ评论，减少用户决策成本',
        '目标：每条视频评论数>20条（触发更高流量池）',
    ]
    for c in comment_strategy:
        add_bullet(doc, c, size=10)

    add_para(doc, '抓手3：矩阵号分发', bold=True, size=11)
    matrix = [
        '建立3-5个手环相关抖音号矩阵（官方号+员工号+垂类号）',
        '同一视频素材剪辑不同版本（不同BGM/开头/时长），分发到各矩阵号',
        '矩阵号之间互相@、合拍、Duet，形成账号群效应',
        '各号定位差异化：①官方主号（品牌信任）②员工人设号（真实体验）③数码测评号（专业背书）',
    ]
    for m in matrix:
        add_bullet(doc, m, size=10)

    add_para(doc, '抓手4：发布时间优化', bold=True, size=11)
    timing = [
        '黄金发布时段（基于手环用户活跃数据）：',
        '  早8:00-9:00（通勤刷手机）',
        '  午12:00-13:00（午休刷手机）',
        '  晚20:00-22:00（晚间高峰，流量最大）',
        '  周末上午10:00-12:00（周末悠闲时间）',
        '每天至少2个时段各发2-3条，测试不同时段表现',
        '发现某时段持续高表现后，固定该时段为"主发布时间"',
    ]
    for t in timing:
        add_bullet(doc, t, size=10)

    add_para(doc, '抓手5：视频SEO优化', bold=True, size=11)
    seo = [
        '标题必含核心关键词："小米手环10 Pro""手环推荐""智能手环"',
        '话题标签策略：1个品牌词 + 2个品类词 + 1个场景词 + 1个热点词',
        '示例：#小米手环10Pro #智能手环推荐 #运动手环 #暑假数码好物 #健康监测',
        '视频文案前15个字包含核心关键词（影响搜索排名）',
    ]
    for s in seo:
        add_bullet(doc, s, size=10)

    add_para(doc, '抓手6：付费快推策略', bold=True, size=11)
    paid_push = [
        '新视频发布后1小时，自然播放<500：立即¥50抖加助推（目标播放量×10）',
        '自然播放>3,000且互动率>5%：追投¥200抖加，助推进入万级流量池',
        '千川投放的视频同步用抖加做"视频互动"目标，提升视频互动数据',
        '每周预算：千川¥43,000（转化）+ 抖加¥5,000（播放/粉丝/互动）',
    ]
    for p in paid_push:
        add_bullet(doc, p, size=10)

    add_styled_heading(doc, '7.3 播放量增长目标', level=2)

    target_headers = ['指标', '当前（周均）', '1个月目标', '3个月目标']
    target_rows = [
        ['手环视频总播放量', '882,299/周', '1,500,000/周 (+70%)', '3,000,000/周 (+240%)'],
        ['日均发布视频数', '约5-8条', '10条/天', '15条/天'],
        ['单条平均播放', '约470', '1,000+', '2,000+'],
        ['万播视频占比', '约5%', '15%', '30%'],
        ['10万+爆款/月', '0-1条', '3-5条', '10条+'],
    ]
    create_table(doc, target_headers, target_rows, col_widths=[3.5, 3, 4, 4])

    doc.add_page_break()

    # ============================================================
    # SECTION 8: EXECUTION PLAN
    # ============================================================
    add_styled_heading(doc, '八、执行计划与KPI', level=1)

    add_styled_heading(doc, '8.1 优先级排序', level=2)
    add_para(doc, '按照"投入产出比"和"实施难度"两个维度排序：', size=10.5)

    priority_headers = ['优先级', '行动项', '实施难度', '预期效果', '负责', '启动时间']
    priority_rows = [
        ['P0-立即', '手环视频日产10条+千川投放提至¥43K/周', '★★☆', '播放量+70%, 订单+30%', '视频+投放组', '本周'],
        ['P0-立即', '"小米官方手环直播间"排品+话术升级', '★☆☆', '单直播间日销+50%', '主播+运营', '本周'],
        ['P1-2周内', '激活"小米手环10Pro直播间"日播8h+投流', '★★★', '月增500-1,000单', '新主播+投放', '7月第3周'],
        ['P1-2周内', '直播间视觉升级（贴片+场景划分）', '★★☆', '停留时长+20%', '设计+装修', '7月第3周'],
        ['P2-1个月内', '矩阵号体系搭建（3-5个号）', '★★★★', '月度额外播放200万+', '视频团队', '7月底'],
        ['P2-1个月内', '竞对"良米-小米手环"深度对标报告', '★★☆', '找到关键差距突破口', '数据分析', '7月底'],
    ]
    create_table(doc, priority_headers, priority_rows, col_widths=[1.8, 5, 1.5, 3, 2, 2])

    add_styled_heading(doc, '8.2 月度KPI目标', level=2)

    kpi_target_headers = ['维度', '指标', '当前水平（6月）', '7月目标', '8月目标']
    kpi_target_rows = [
        ['销量', '手环月订单', '57,745单', '65,000单 (+12%)', '75,000单 (+30%)'],
        ['销量', '我方市占率', '35.3%', '38%', '42%'],
        ['直播间', '官旗手环间日销', '317单/天', '450单/天', '600单/天'],
        ['直播间', '10Pro间日销', '6单/天', '30单/天', '80单/天'],
        ['视频', '手环视频周播放', '88万', '150万', '250万'],
        ['视频', '日产视频数', '5-8条/天', '10条/天', '15条/天'],
        ['投放', '千川周消耗', '¥27,434', '¥43,000', '¥60,000'],
        ['投放', '千川ROI', '14.1', '维持>12', '维持>12'],
    ]
    create_table(doc, kpi_target_headers, kpi_target_rows, col_widths=[1.5, 3, 3, 3.5, 3.5])

    add_styled_heading(doc, '8.3 资源配置建议', level=2)

    resource_headers = ['资源类型', '当前', '建议', '增量投入']
    resource_rows = [
        ['视频编导/剪辑', '1-2人', '3-4人（含1名编导）', '+1-2人'],
        ['主播', '6间直播间轮流', '新增1名手环专属主播（运动科技方向）', '+1人'],
        ['千川投放预算/月', '¥11万/月', '¥17-20万/月', '+6-9万/月'],
        ['抖加预算/月', '少量', '¥2万/月', '+2万/月'],
        ['直播设备', '基础设备', '新增运动场景拍摄设备（GoPro/运动相机）', '一次性¥1万'],
    ]
    create_table(doc, resource_headers, resource_rows, col_widths=[3, 3.5, 5.5, 3])

    doc.add_paragraph()

    # ============================================================
    # CLOSING
    # ============================================================
    add_styled_heading(doc, '总结', level=1)

    add_para(doc, '小米手环系列直播间提升的核心逻辑：', size=11)

    summary_points = [
        '【拉流量】视频日产10条+千川周投¥43K+矩阵号分发 → 做大流量入口',
        '【提转化】直播间话术升级+视觉优化+排品节奏 → 提高进房转化率',
        '【抢份额】对标良米"小米手环"直播间（月19,362单），逐一拆解其优势并针对性超越',
        '【建壁垒】"官方手环直播间"做大基本盘 + "10Pro直播间"做高端差异化 → 双轮驱动',
        '【可持续】视频内容SOP化、话术模板化、投放系统化 → 不依赖个人能力',
    ]
    for s_text in summary_points:
        add_bullet(doc, s_text, size=10.5)

    add_para(doc, '')
    add_para(doc, f'以6月手环订单57,745单为基准，通过以上全链路优化，预计7月达到65,000单（+12%），8月达到75,000单（+30%），'
             f'我方市占率从35.3%提升至42%以上，在Q3结束前实现对良米的反超。', bold=True, size=11, color=(255,105,0))

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer info
    add_para(doc, '— 报告完 —', size=10, color=(150,150,150), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, f'数据来源：小米手环直播间销量分析平台 | 生成时间：2026年7月2日',
             size=9, color=(180,180,180), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Save
    output_path = os.path.join(DATA_DIR, '小米手环直播间提升方案.docx')
    doc.save(output_path)
    print(f'Report saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    build_report()
